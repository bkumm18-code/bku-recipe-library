import argparse, json
from datetime import date, timedelta
from pathlib import Path
from docx import Document

SUBJECTS=['総合日本語','聴解','読解','文字','発音','会話','作文','生活一般に関する知識','円滑な技能等の修得等に資する知識']
FIXED={(1,1),(2,12),(3,2),(7,19),(12,25)}
def main():
    p=argparse.ArgumentParser(); p.add_argument('--template',required=True); p.add_argument('--start',required=True); p.add_argument('--end',required=True); p.add_argument('--visa',default='kaigo'); p.add_argument('--holidays'); p.add_argument('--output',required=True); p.add_argument('--instructor',default='Christi'); p.add_argument('--location',default='BKU Co., Ltd'); p.add_argument('--student'); a=p.parse_args()
    st=date.fromisoformat(a.start); en=date.fromisoformat(a.end)
    if en<st: raise SystemExit('end date must not precede start date')
    hol=set(json.loads(Path(a.holidays).read_text(encoding='utf-8'))) if a.holidays else set(); days=[]; d=st
    while d<=en:
        if d.weekday()<5 and d.isoformat() not in hol and (d.month,d.day) not in FIXED: days.append(d)
        d+=timedelta(days=1)
    subs=list(SUBJECTS)
    if a.visa.lower() in ('kaigo','caregiving','介護'): subs.insert(7,'介護の日本語')
    q,r=divmod(len(days),len(subs)); assigned=[s for i,s in enumerate(subs) for _ in range(q+(i<r))]
    doc=Document(a.template)
    for para in doc.paragraphs:
        if '実施期間' in para.text:
            para.text=f'（実施期間　　　{st.year}年　 {st.month:02d}月　 {st.day:02d}日から　　　{en.year}年　 {en.month:02d}月　 {en.day:02d}日まで）'; break
    t=doc.tables[0]
    while len(t.rows)>1: t._tbl.remove(t.rows[-1]._tr)
    for d,s in zip(days,assigned):
        vals=[d.strftime('%Y/%m/%d'),'9:00 ～ 17:00',s,a.instructor,a.location,'']; cells=t.add_row().cells
        for c,v in zip(cells,vals): c.text=v
    if a.student and len(doc.tables)>1 and len(doc.tables[1].rows)>1:
        row=doc.tables[1].rows[1].cells; row[0].text='1'; row[1].text=a.student; row[2].text=''; row[3].text=''
    out=Path(a.output); out.parent.mkdir(parents=True,exist_ok=True); doc.save(out)
    print(json.dumps({'output':str(out),'training_days':len(days),'excluded_holidays':sorted(hol)},ensure_ascii=False))
if __name__=='__main__': main()

